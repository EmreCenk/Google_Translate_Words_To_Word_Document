
from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
import docx
from datetime import datetime
from time import sleep



class wordhelper:
    def __init__(self):


        self.browser = webdriver.Chrome(r"chromedriver.exe")
        #self.browser.minimize_window()
        self.thingtodo = ActionChains(self.browser)

    def urltolistword(self, url):
        x = url
        x+=" "
        x = x[x.find("text") + 5:-1]

        x = x.replace("%20", " ").replace("%0A", "\n").replace("%3A",":").replace("%2",",")
        x = x.split("\n")
        inx=[]
        for i in x:
            if i not in inx and i not in [""," ", "  ","    "] and len(i)<30:
                inx.append(i)
        return inx

    def getwords(self, url):

        self.browser.get(url)
        leftside=self.urltolistword(url)
        print("left side is:")
        print(leftside)
        possible=["J0lOec","translation"]


        try:
            bimo = self.browser.find_element_by_class_name("translation")
        except:
            bimo="\n"

        try:
            curbim=len(bimo.text.split("\n"))
        except:
            curbim=0

        n=0
        while curbim<len(leftside):
            try:
                bimo = self.browser.find_element_by_class_name(possible[n%len(possible)])
            except:
                pass
            try:
                curbim = len(bimo.text.split("\n"))
            except:
                curbim = 0

            n+=1
        translation = bimo.text.split("\n")
        needed={}

        for i in range(0, len(translation)):
            needed[leftside[i]] = translation[i]

        print("'needed' is:")
        print(needed)

        return needed

    def writetoworddocument(self, path, dictio):
        


        mydoc = docx.Document(path)
        date=datetime.now().strftime("%B %d, %Y %H:%M")
        fer=mydoc.add_paragraph()
        fer.add_run("On " + date + ", the following words were recorded: ").bold=True
        for i in dictio:
            x = i + ": " + dictio[i]
            mydoc.add_paragraph(x)
            print(x)

        mydoc.save(path)
        print("The words have been written to the document.")
        print("\nThe document has been saved.\n")


    def overall(self,url,path):
        dictio=self.getwords(url)
        print("\nWriting to document... This may take a few seconds.\n")

        self.browser.close()


        self.writetoworddocument(path=path,dictio=dictio)
        print("\n")




input("Paste the link. After doing so, press enter to start program. You will notice a google chrome tab open "
"once the program starts. \n")

self=wordhelper()
print(self.getwords(url=url))
path="Enter a path to a word document"
self.overall(url=url, path=path)
